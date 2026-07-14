#!/usr/bin/env python3
"""Surgical Section-11 failure-PR repair: switch the failure PR numerator from the scenario-local
path-library optimum to the strict full-MCF optimum (same formula as normal PR: min(1, strict/achieved)).

ONLY Section-11 failure-PR cells and the Section-11 reference wording are changed in the governing report
(RG_GNN_LPD_FIX1_Final_Report_..._MLU_UPGRADED.docx). Everything else (Sections 1-10, 12, 13, 14-16, normal
N=3976, FlexDATE, SDN Section 13) is left byte-identical because we edit table cells in place (no rebuild).
Achieved MLU, DB, decision ms, disconnected counts, actions, N are all preserved unchanged.

Aborts if the strict-failure cache is not complete for every (topology, scenario, tm) in the failure per-cycle
artifact. Also updates the source CSVs' PR columns + provenance for reproducibility (achieved etc. unchanged).
Safety-mode artifacts are not used.
"""
import sys, shutil
import numpy as np, pandas as pd
from pathlib import Path
sys.path.insert(0, "/Users/moahaimentalib/Desktop/f_flex_network_code_clean")
from docx import Document

RC = Path("/Users/moahaimentalib/Desktop/f_flex_network_code_clean/results/gnn_lpd_dqn_selective_db_lp/condition_compliant_k10_k50")
CM = RC / "FINAL_REPORT_FIX1" / "completed_metrics"
SF = RC / "STRICT_FAILURE_FULL_MCF_PR"
RPT = RC / "FINAL_REPORT_FIX1"
DOCX = RPT / "RG_GNN_LPD_FIX1_Final_Report_200VTL_N3976_FINAL_METHOD_AUDITED_PRREF_STRICTALL_MLU_UPGRADED.docx"
FAILURE_METHODS = ["ECMP","OSPF-weighted shortest-path routing","Top-K Demand (K50)","Bottleneck Top-K (K50)","GNN-only fixed K50","GNN+LPD fixed K50","Final RG-GNN-LPD"]
SCEN = ["single_link_failure","two_link_failure","three_link_failure","capacity_degradation_50","spike","mixed_spike_failure"]
DISP_T = {"abilene":"abilene","geant":"geant","cernet":"cernet","sprintlink":"sprintlink","tiscali":"tiscali","ebone":"ebone","germany50":"germany50","vtlwavenet2011":"vtlwavenet2011"}
def pct(v): return "N/A" if pd.isna(v) else f"{float(v):.3f}%"

# ---- load per-cycle + strict cache ----
pc = pd.read_csv(CM / "failure_baseline_comparison_fix1_per_cycle.csv")
strict = {}
for t in ["abilene","geant","ebone","cernet","germany50","sprintlink","tiscali","vtlwavenet2011"]:
    p = SF / "_partial" / f"{t}.progress.csv"
    if p.exists():
        for r in pd.read_csv(p).itertuples():
            strict[(t, r.scenario, int(r.tm_id))] = float(r.strict_full_mcf_MLU)

# ---- completeness gate ----
missing = [(r.Topology, r.Scenario, int(r.tm_index)) for r in pc.itertuples()
           if (r.Topology, r.Scenario, int(r.tm_index)) not in strict]
if missing:
    print(f"ABORT: strict cache incomplete — {len(set(missing))} unique states missing, e.g. {sorted(set(missing))[:3]}")
    sys.exit(1)
print(f"[ok] strict cache complete for all {pc[['Topology','Scenario','tm_index']].drop_duplicates().shape[0]} states")

# ---- recompute new PR (plain min(1, strict/achieved), same as normal) ----
pc["strict_full_mcf_MLU"] = [strict[(r.Topology, r.Scenario, int(r.tm_index))] for r in pc.itertuples()]
pc["PR_old"] = pc["PR"]
pc["PR"] = np.minimum(1.0, pc["strict_full_mcf_MLU"] / pc["achieved_mlu"].where(pc["achieved_mlu"] > 0, np.nan)).fillna(0.0)
pc["PR>=0.90 flag"] = (pc["PR"] >= 0.90).astype(int)
pc["PR>=0.95 flag"] = (pc["PR"] >= 0.95).astype(int)
pc["pr_reference_type"] = "strict_full_mcf_exact_scenario_state"

def agg_sm(df):
    return df.groupby(["Scenario","Method"]).apply(
        lambda g: pd.Series(dict(MeanPR=g.PR.mean()*100, pr90=(g.PR>=0.90).mean()*100, pr95=(g.PR>=0.95).mean()*100)),
        include_groups=False)
NEW = agg_sm(pc); OLD = pc.assign(PR=pc.PR_old).pipe(agg_sm)   # OLD for table identification
fin = pc[pc.Method == "Final RG-GNN-LPD"]
NEWt = fin.groupby(["Scenario","Topology"]).apply(lambda g: pd.Series(dict(MeanPR=g.PR.mean()*100, pr90=(g.PR>=0.90).mean()*100)), include_groups=False)
OLDt = fin.assign(PR=fin.PR_old).groupby(["Scenario","Topology"]).apply(lambda g: pd.Series(dict(MeanPR=g.PR.mean()*100)), include_groups=False)

# ---- backup + open docx ----
shutil.copy(DOCX, DOCX.with_suffix(".prestrict_backup.docx"))
doc = Document(str(DOCX))
def cellf(x):
    try: return float(str(x).replace("%","").strip())
    except: return None
def set_cell(cell, text):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]: r.text = ""
    else:
        cell.text = text

edits = 0
# --- T19-T24: the 6 per-scenario Method tables appear in document order = FAILURE_SCENARIO_ORDER ---
method_tables = []
for tb in doc.tables:
    hdr = [c.text.strip() for c in tb.rows[0].cells]
    if hdr[0] == "Method" and not any("PR(fail)" in c for c in hdr) and any("Disconnected" in c for c in hdr) and len(tb.rows) == 8:
        if [tb.rows[i].cells[0].text.strip() for i in range(1, 8)] == FAILURE_METHODS:
            method_tables.append(tb)
assert len(method_tables) == len(SCEN), f"expected {len(SCEN)} method tables, found {len(method_tables)}"
# sanity: disconnected-OD column (unchanged) must match the CSV for the assigned scenario/ECMP
disc_col = next(i for i, c in enumerate([c.text.strip() for c in method_tables[0].rows[0].cells]) if "Disconnected" in c)
for scen, tb in zip(SCEN, method_tables):
    ecmp_disc_doc = cellf(tb.rows[1].cells[disc_col].text)
    ecmp_disc_csv = float(pc[(pc.Scenario == scen) & (pc.Method == "ECMP")]["Disconnected ODs"].max())
    assert abs((ecmp_disc_doc or 0) - ecmp_disc_csv) < 0.5, f"order mismatch at {scen}: doc disc={ecmp_disc_doc} csv={ecmp_disc_csv}"
    for j, m in enumerate(FAILURE_METHODS):
        row = tb.rows[1+j]; nv = NEW.loc[(scen, m)]
        for ci, val in [(1, nv.MeanPR), (2, nv.pr90), (3, nv.pr95)]:
            set_cell(row.cells[ci], pct(val))
        edits += 1
    print(f"[T19-24] {scen}: updated {len(FAILURE_METHODS)} method rows (disc sanity ok)")

# --- T25: Scenario|Topology|Mean PR|PR>=0.90|PR(fail)|... (49 rows) ---
for tb in doc.tables:
    hdr = [c.text.strip() for c in tb.rows[0].cells]
    if not (hdr[0] == "Scenario" and hdr[1] == "Topology" and any("PR(fail)" in c for c in hdr)):
        continue
    for i in range(1, len(tb.rows)):
        row = tb.rows[i]; sc = row.cells[0].text.strip(); tp = row.cells[1].text.strip()
        if (sc, tp) not in NEWt.index: continue
        nv = NEWt.loc[(sc, tp)]
        oldpf = row.cells[4].text.strip()
        vals = {2: nv.MeanPR, 3: nv.pr90}
        vals[4] = nv.MeanPR if oldpf not in ("N/A", "", "-") else np.nan  # PR(fail)=MeanPR for failures; keep N/A for spike
        for ci, val in vals.items():
            new_txt = pct(val) if not (ci == 4 and pd.isna(val)) else oldpf
            set_cell(row.cells[ci], new_txt)
        edits += 1
    print(f"[T25] updated per-topology robustness rows")

# --- Section-11 reference wording ---
NEW_NOTE = ("PR-reference note: Section-11 failure PR uses the strict full-MCF optimum recomputed for the exact "
            "scenario state (failed links, capacity modifications, and/or demand transformation applied first), "
            "the same strict full-MCF numerator used by the normal N=3976 protocol. The earlier scenario-local "
            "path-library optimum is no longer used as the failure PR numerator; the strict-failure reference cache "
            "is STRICT_FAILURE_FULL_MCF_PR/.")
wchg = 0
for para in doc.paragraphs:
    if "scenario-local path-library optimum" in para.text:
        for r in para.runs: r.text = ""
        para.runs[0].text = NEW_NOTE if para.runs else None
        if not para.runs: para.add_run(NEW_NOTE)
        wchg += 1
    elif "the same scenario-local PR reference" in para.text:
        para.text = para.text.replace("the same scenario-local PR reference", "the same strict full-MCF PR reference (exact scenario state)")
        wchg += 1
print(f"[wording] updated {wchg} paragraph(s)")

doc.save(str(DOCX))
# persist recomputed CSVs (PR + provenance changed; achieved/DB/ms/disc unchanged)
pc.drop(columns=["PR_old"]).to_csv(CM / "failure_baseline_comparison_fix1_per_cycle.csv", index=False)
print(f"[done] edits={edits}, docx saved. Backup: {DOCX.with_suffix('.prestrict_backup.docx').name}")
