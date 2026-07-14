#!/usr/bin/env python3
"""Surgical Section-11 insertion ONLY: append (before Section 12) the Full-Stream Periodic Single-Link Failure
Results table (strict full-MCF, failure-cycle metrics), the Failure Action Counts by Topology/Scenario table,
and a reference to the per-TM/per-cycle failure action-trace artifact. Nothing else is changed."""
import numpy as np, pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RC = Path("/Users/moahaimentalib/Desktop/f_flex_network_code_clean/results/gnn_lpd_dqn_selective_db_lp/condition_compliant_k10_k50")
RPT = RC / "FINAL_REPORT_FIX1"
R = RPT / "RG_GNN_LPD_FIX1_Final_Report_200VTL_N3976_FINAL_METHOD_AUDITED_PRREF_STRICTALL_MLU_UPGRADED.docx"
SF = RC / "STRICT_FAILURE_FULL_MCF_PR"
HDR = "2E5496"; DISP = {"abilene":"Abilene","geant":"GÉANT","cernet":"CERNET","sprintlink":"Sprintlink","tiscali":"Tiscali","ebone":"Ebone","germany50":"Germany50","vtlwavenet2011":"VtlWavenet2011"}
SC_ORDER = ["single_link_failure","two_link_failure","three_link_failure","capacity_degradation_50","spike","mixed_spike_failure"]
SC_DISP = {"single_link_failure":"single-link","two_link_failure":"two-link","three_link_failure":"three-link","capacity_degradation_50":"capacity-50%","spike":"spike","mixed_spike_failure":"mixed spike+failure"}
ACTS = ["KEEP","K50","K100","K200","K300","K500","K800"]
TOPO = ["abilene","geant","cernet","sprintlink","tiscali","ebone","germany50","vtlwavenet2011"]

import shutil; shutil.copy(R, R.with_suffix(".pre_s11tables_backup.docx"))
doc = Document(str(R))

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), fill); tcPr.append(sh)
def setf(run, size, bold=False, white=False, italic=False):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    if white: run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# find Section-12 anchor
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith("12.") and "Zero-Shot" in p.text: anchor = p; break
assert anchor is not None, "Section 12 anchor not found"

def before(el): anchor._p.addprevious(el)
def add_heading(text):
    p = doc.add_paragraph(); r = p.add_run(text); setf(r, 10, bold=True); before(p._p)
def add_note(text):
    p = doc.add_paragraph(); r = p.add_run(text); setf(r, 8.5, italic=True); before(p._p)
def add_table(headers, rows, fs=8.0, first_col=1.15, total=9.6):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = str(h); shade(c, HDR)
        for p0 in c.paragraphs:
            for r in p0.runs: setf(r, fs, bold=True, white=True)
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = str(v)
            for p0 in cs[i].paragraphs:
                for r in p0.runs: setf(r, fs)
    n = len(headers); rest = max((total - first_col) / max(n - 1, 1), 0.55); widths = [first_col] + [rest] * (n - 1)
    tw = [int(round(w * 1440)) for w in widths]
    t.autofit = False; t.allow_autofit = False
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); t._tbl.tblPr.append(lay)
    grid = t._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), tw): gc.set(qn("w:w"), str(w))
    for row in t.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i]); tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")): tcPr.remove(old)
            e = OxmlElement("w:tcW"); e.set(qn("w:w"), str(tw[i])); e.set(qn("w:type"), "dxa"); tcPr.append(e)
    before(t._tbl)

# ---- 1) Periodic single-link failure results (strict full-MCF, failure-cycle metrics) ----
per = pd.read_csv(SF / "periodic_single_link_STRICT_failurecycle_table.csv").set_index("Topology")
add_heading("Full-Stream Periodic Single-Link Failure Results (strict full-MCF numerator)")
prows = []
name_by_disp = {DISP[t]: t for t in TOPO}
for _, r in per.iterrows():
    prows.append([r.name, int(r.TMs), int(r.n_fail), r.every, f"{r.Mean_PR*100:.3f}%",
                  f"{r.Fail_Min_PR*100:.3f}%", f"{r.Fail_Mean_PR*100:.3f}%", f"{r.Fail_PR_ge90:.1f}%"])
add_table(["Topology","TMs","Failure cycles","Interval","Overall Mean PR","Failure-Cycle Min PR","Failure-Cycle Mean PR","Failure-Cycle PR>=0.90"], prows, fs=8.0, first_col=1.3)
add_note("PR numerator is the strict full-MCF optimum recomputed for each exact failure state (same reference as the "
         "normal protocol). Failure-cycle columns isolate the transient-failure cycles; the overall-stream mean also "
         "includes the intervening normal cycles (whose per-cycle behavior is reported in the normal table). Failure "
         "cycles are transient single-link failures rotating over the highest-capacity links.")

# ---- 2) Failure action counts by topology and scenario ----
pc = pd.read_csv(RC / "FINAL_REPORT_FIX1" / "completed_metrics" / "failure_baseline_comparison_fix1_per_cycle.csv")
fin = pc[pc.Method == "Final RG-GNN-LPD"]
add_heading("Failure Action Counts by Topology and Scenario (Final RG-GNN-LPD, 20 TMs per block)")
arows = []
for t in TOPO:
    for sc in SC_ORDER:
        g = fin[(fin.Topology == t) & (fin.Scenario == sc)]; vc = g.action.value_counts()
        most = vc.idxmax() if len(vc) else "-"
        arows.append([DISP[t], SC_DISP[sc]] + [int(vc.get(a, 0)) for a in ACTS] + [most])
add_table(["Topology","Scenario"] + ACTS + ["Most used"], arows, fs=7.5, first_col=1.25)
add_note("Per-cycle DDQN action selections aggregated over the 20-TM failure block. The complete per-TM/per-cycle "
         "action trace for all topologies and scenarios is provided as the artifact "
         "STRICT_FAILURE_FULL_MCF_PR/failure_action_trace_per_cycle.csv (960 rows: topology, scenario, tm_index, "
         "action, selected_K, disconnected ODs).")

doc.save(str(R))
print(f"[done] inserted 2 tables + notes before Section 12. Backup: {R.with_suffix('.pre_s11tables_backup.docx').name}")
