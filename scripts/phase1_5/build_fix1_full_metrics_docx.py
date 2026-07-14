#!/usr/bin/env python3
"""Build the FIX1 FULL-METRICS DOCX (RG-GNN-LPD) from fix1_full_metrics.json + figures + source CSVs.
Honest labels for every gap. Landscape, wide metric tables. Does NOT overwrite short/comprehensive reports."""
import json
from pathlib import Path
import numpy as np, pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RC = Path("/Users/moahaimentalib/Desktop/f_flex_network_code_clean/results/gnn_lpd_dqn_selective_db_lp/condition_compliant_k10_k50")
SDN = RC.parent / "sdn_mininet_clean"
F1 = RC / "FULLDATA_GATED_PRESERVED_FIX1"; FAIL = RC / "PERIODIC_FAILURE_FIX1"
RPT = RC / "FINAL_REPORT_FIX1"; FIG = RPT / "figures"
LABEL = "RG-GNN-LPD"
TOPO = ["abilene","geant","cernet","sprintlink","tiscali","ebone","germany50","vtlwavenet2011"]
DISP = {"abilene":"Abilene","geant":"GEANT","cernet":"CERNET","sprintlink":"Sprintlink","tiscali":"Tiscali","ebone":"Ebone","germany50":"Germany50","vtlwavenet2011":"VtlWavenet2011"}
ZS = {"germany50","vtlwavenet2011"}
SRC = {"abilene":"Real / SNDlib","geant":"Real / SNDlib","cernet":"MGM (synthetic-real)","sprintlink":"RocketFuel / MGM","tiscali":"RocketFuel / MGM","ebone":"RocketFuel / MGM","germany50":"Real / SNDlib","vtlwavenet2011":"Topology Zoo / MGM"}
TRAIN = {"abilene":2016,"geant":672,"cernet":200,"sprintlink":200,"tiscali":200,"ebone":200,"germany50":0,"vtlwavenet2011":0}
TESTN = {"abilene":2016,"geant":672,"cernet":200,"sprintlink":200,"tiscali":200,"ebone":200,"germany50":288,"vtlwavenet2011":40}
GNN_MS = {"abilene":1,"geant":7,"cernet":7,"sprintlink":10,"tiscali":10,"ebone":4,"germany50":26,"vtlwavenet2011":26}
ACTS = ["KEEP","K50","K100","K200","K300","K500","K800"]
NFAIL = {'abilene':12,'geant':4,'cernet':2,'sprintlink':2,'tiscali':2,'ebone':2,'germany50':4,'vtlwavenet2011':2}
M = json.load(open(RPT/"fix1_full_metrics.json"))
adist = pd.read_csv(F1/"action_distribution.csv").set_index("Topology")
ev = pd.read_csv(F1/"fulldata_gated_eval_per_cycle.csv")
flex = pd.read_csv(F1/"flexdate_comparison.csv").set_index("topology")
sdn = pd.read_csv(SDN/"sdn_summary.csv")
HDR="2E5496"

doc=Document(); s=doc.sections[0]; s.orientation=WD_ORIENT.LANDSCAPE; s.page_width=Inches(11); s.page_height=Inches(8.5)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(s,m,Inches(0.7))
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10)
def shade(c,f): tcPr=c._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),f); tcPr.append(sh)
def setf(c,sz,bold=False,white=False):
    for p in c.paragraphs:
        for r in p.runs: r.font.size=Pt(sz); r.font.bold=bold
        if white:
            for r in p.runs: r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
TOTW = 9.6  # landscape content width (inches)
def _fixed_widths(t, c0):
    """Fixed table layout with a wider first column so long topology names never wrap mid-word.
    LibreOffice honors the tblGrid <w:gridCol> widths, so set those (in twips) plus fixed layout."""
    n=len(t.columns); rest=(TOTW-c0)/max(n-1,1); widths=[c0]+[rest]*(n-1)
    tw=[int(round(w*1440)) for w in widths]
    t.autofit=False; t.allow_autofit=False
    tblPr=t._tbl.tblPr; lay=OxmlElement("w:tblLayout"); lay.set(qn("w:type"),"fixed"); tblPr.append(lay)
    # set the table grid columns (what LibreOffice uses for fixed layout)
    grid=t._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        cols=grid.findall(qn("w:gridCol"))
        for gc,w in zip(cols,tw): gc.set(qn("w:w"),str(w))
    for row in t.rows:
        for j,cell in enumerate(row.cells):
            cell.width=Inches(widths[j])
            tcPr=cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")): tcPr.remove(old)
            tcW=OxmlElement("w:tcW"); tcW.set(qn("w:w"),str(tw[j])); tcW.set(qn("w:type"),"dxa"); tcPr.append(tcW)
def table(headers,rows,fs=8,c0=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=str(h); shade(c,HDR); setf(c,fs,True,True)
    for row in rows:
        cs=t.add_row().cells
        for j,v in enumerate(row): cs[j].text=str(v); setf(cs[j],fs)
    if c0: _fixed_widths(t, c0)
    return t
def para(txt,size=10,bold=False,italic=False,space=6,keep_next=False):
    p=doc.add_paragraph(); r=p.add_run(txt); r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic; p.paragraph_format.space_after=Pt(space); p.paragraph_format.keep_with_next=keep_next; return p
def note(txt): para(txt,size=8.5,italic=True,space=8)
def bullets(items,size=9):
    for b in items:
        p=doc.add_paragraph(b); p.paragraph_format.left_indent=Inches(0.3); p.paragraph_format.space_after=Pt(2)
        for r in p.runs: r.font.size=Pt(size)
def h1(t): doc.add_heading(t,level=1)
def fig(name,w=7.2):
    f=FIG/name
    if f.exists(): doc.add_picture(str(f),width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

# Title
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("RG-GNN-LPD Final Report"); r.font.size=Pt(19); r.font.bold=True
q=doc.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=q.add_run("Proposed method: Reward-Gated GNN-LPD Traffic Engineering (RG-GNN-LPD)\nFormal implementation: Full-Data-Trained Topology-Agnostic Bottleneck-Ranking DDQN with Gated First-Cycle Initialization"); r.font.size=Pt(11); r.font.italic=True

# 1
h1("1. Executive Summary")
para(f"The proposed method is Reward-Gated GNN-LPD Traffic Engineering ({LABEL}). It is the full-data-trained "
     "FIX1 learned controller: a trained LP-distilled GraphSAGE criticality scorer (the runtime "
     "`GNNLPDScorer`) produces per-OD scores, a bottleneck-relief ranking term is combined with those scores to "
     "rank critical ODs, a reward-gated Double-DQN policy selects the action (KEEP or a fixed-K "
     "optimize budget); a selected-flow LP computes a feasible routing for the selected ODs; and all nonselected "
     "ODs remain on ECMP. The model is initialized from the previous frozen policy and fine-tuned on the full "
     "training partitions with behavior-preservation distillation; first-TM teacher-action preservation is part of "
     "the disclosed gated initialization mechanism, with a cycle-0 KEEP guard that masks any frozen-teacher KEEP and "
     "forces the best non-KEEP optimize action because no previously accepted routing exists yet. Cycles after the "
     "first TM use the fine-tuned DDQN policy.")
note("Method identity rule for this report: the proposed method name is Reward-Gated GNN-LPD Traffic Engineering "
     "(RG-GNN-LPD). The longer DDQN phrase is the formal implementation name, not a competing final method name. "
     "A separate safety-hardened deployment artifact exists in the repository, but it does not clear the all-topology "
     "MinPR>=0.90 acceptance rule, so this report remains anchored to the normal FIX1 RG-GNN-LPD results and does "
     "not claim an all-cycle PR>0.90 operating mode.")
bullets([
 f"Final method = {LABEL} (FIX1), full-data trained.",
 "8-topology normal results: all PR>=0.90 (Section 4).",
 "FlexDATE = 3/5 learned wins (Abilene, CERNET, GEANT).",
 "4/5 only with a separate deployable Sprintlink K800 route (PR 0.9993, p95 ~325 ms); Sprintlink is NOT a learned win.",
 "Failure: periodic single-link robustness stress-test (Section 12), not a normal-runtime claim.",
 "SDN/Mininet: a retained operational validation artifact (Abilene+GEANT), NOT rerun on FIX1 (Section 11).",
 "Runtime on the largest topologies (VtlWavenet, Tiscali) is borderline-near-500 ms; not a strict all-repetition guarantee."])

# 2
h1("2. Method / Algorithm Pipeline")
para(f"{LABEL} pipeline (per traffic-matrix cycle):", bold=True, size=10)
table(["Stage","Component","Role"],
 [["1. Input","Traffic matrix (TM)","Per-cycle demand vector over OD pairs"],
  ["2. Criticality scorer","LP-distilled GraphSAGE (`GNNLPDScorer`)","Per-OD criticality score from graph + load features, trained with LP-distilled supervision"],
  ["3. Bottleneck term","Relief score","Topology-agnostic bottleneck-relief signal used to prioritize WHICH ODs matter most"],
  ["4. Ranking","Bottleneck + GNN-LPD","OD-criticality ranking = relief + 0.3*GNN-LPD score (selects WHICH ODs)"],
  ["5. Policy","Reward-gated DDQN (the gate)","Selects the ACTION: KEEP or a fixed-K optimize budget (K50..K800)"],
  ["6. Action","Selected K / KEEP","Number of top-ranked ODs to re-route this cycle (or keep previous routing)"],
  ["7. Optimizer","Selected-flow LP (db-budgeted)","Feasible routing for the selected ODs under a disturbance budget"],
  ["8. Background","ECMP","Routing for all nonselected ODs (consistent ECMP background)"],
  ["9. Carry-forward","Accepted splits","Routing carried to the next cycle (KEEP reuses it; DB measured vs it)"],
  ["10. Init","First-TM teacher-action preservation","Cycle-0 action from the frozen teacher (disclosed gated init), with cycle-0 KEEP masked to the best non-KEEP action"],
  ["11. Gate rule","Gated first-cycle","After the cycle-0 KEEP guard, optimize with K>=300 opens db_budget=1.0 for cold-start convergence"]], fs=8.5)
note("The runtime artifact exposes one trained `GNNLPDScorer`: a GraphSAGE-style OD scorer trained with LP-distilled "
     "supervision. The FIX1 runtime does not log a separate per-cycle LPD column; instead, the LP-distilled teacher "
     "signal is embedded in the trained scorer itself. Reward gate / DDQN = action policy (KEEP / K budget). "
     "Selected-flow LP = feasible routing optimizer. ECMP = routing for nonselected ODs. Cycles after the first TM "
     "use the fine-tuned DDQN policy. The cold-start guard is explicit: if the frozen teacher predicts KEEP at TM0, "
     "KEEP is masked and the best non-KEEP optimize action is used instead.")

# 3
h1("3. Dataset / Train / Validation / Test")
rows=[]
for t in TOPO:
    tn="40 (normal); 200 (failure stress)" if t=="vtlwavenet2011" else str(TESTN[t])
    rows.append([DISP[t],SRC[t],"zero-shot" if t in ZS else "seen",str(TRAIN[t]),"see note",tn,"Yes" if t in ZS else "No"])
table(["Topology","Source","Role","Train cycles","Val","Test cycles","Zero-shot?"], rows, fs=8.5)
note("Validation: no separate held-out validation stream was carved out per topology. Checkpoint/model selection "
     "used a validation-selection log over the fine-tune epoch checkpoints (validation_selection_log.csv) rather "
     "than a distinct held-out cycle count, so an exact Val cycle count is not defined ('see note'); it is not "
     "invented here. Germany50 and VtlWavenet are zero-shot (train=0, val=0). No test leakage.")

# 4
h1("4. Normal-Scenario Metrics (FIX1)")
rows=[]
for t in TOPO:
    m=M[t]
    rows.append([DISP[t],m["nod"],m["N"],f"{m['meanPR']:.4f}",f"{m['medPR']:.4f}",f"{m['pr99']:.0f}%",f"{m['pr95']:.0f}%",
        f"{m['pr90']:.0f}%",f"{m['minPR']:.4f}",f"{m['meanDB']:.4f}",f"{m['p95DB']:.4f}",
        f"{m['mean_ms']:.1f}",f"{m['p95_ms']:.1f}",f"{m['max_ms']:.1f}",m["most_action"]])
table(["Topology","OD pairs","N","MeanPR","MedPR","PR>=.99","PR>=.95","PR>=.90","MinPR","MeanDB","P95DB","mean ms","p95 ms","max ms","Top action"], rows, fs=7.5, c0=1.05)
note("PR percentiles, decision-time (mean/p95/max), and Top action are computed per-cycle from "
     "fulldata_gated_eval_per_cycle.csv. MeanDB / P95DB are the realized aggregate disturbances "
     "(normal_8topo_vs_target.csv); per-cycle realized DB is NOT logged (the per-cycle DB column stores the budget), "
     "so Median DB, Max DB, and a DB CDF are NOT AVAILABLE in the current artifact (see MISSING_BASELINES.md). "
     "MLU columns are intentionally omitted: raw MLU is in per-topology capacity units (Abilene ~0.05 ... VtlWavenet "
     "~1.8e4) and is NOT normalized across topologies, so absolute MLU is not comparable; PR (= optimum/achieved) is "
     "the normalized routing-quality metric and is reported instead. MLU/opt and MLU/ECMP are not reported because "
     "the eval PR uses the strict full-MCF numerator while the cached optimum/ECMP series use a different basis.")
note("Timing caveat: a separate clean-timing batch (fix1_clean_timing_per_cycle.csv) places Tiscali p95 at 506.3 ms "
     "and VtlWavenet near the boundary; do NOT claim every repetition is strictly < 500 ms. Use 'borderline-near-500' "
     "for VtlWavenet/Tiscali. Smaller topologies are comfortably under budget.")
para("Dedicated FlexDATE overlapping-topologies comparison (FIX1 learned model):", bold=True, size=9.5)
rows=[]
for t in ["abilene","cernet","geant","sprintlink","tiscali"]:
    r = flex.loc[t]
    pr_win = "Yes" if float(r.our_PR) >= float(r.target_PR) else "No"
    db_win = "Yes" if float(r.our_DB) < float(r.target_DB) else "No"
    rows.append([
        DISP[t],
        f"{float(r.target_PR)*100:.2f}%",
        f"{float(r.our_PR)*100:.2f}%",
        pr_win,
        f"{float(r.target_DB)*100:.2f}%",
        f"{float(r.our_DB)*100:.2f}%",
        db_win,
        "Yes" if bool(r.win) else "No",
        str(r.note),
    ])
table(["Topology","FlexDATE PR","Our PR","PR win","FlexDATE DB","Our DB","DB win","Learned overall","Reference"], rows, fs=8)
note("This is the dedicated learned overlap table sourced from FULLDATA_GATED_PRESERVED_FIX1/flexdate_comparison.csv. "
     "It shows 3/5 learned wins (Abilene, CERNET, GEANT). Sprintlink is not counted as a learned win in FIX1 "
     "(learned PR 99.61% < target 99.90%), even though the report separately notes a deployable Sprintlink K800 route "
     "outside the learned FIX1 comparison. Tiscali remains non-winning and its reference row is marked informal.")

# 5
h1("5. Baseline Comparison")
para("Real FIX1 reference points (per-cycle, from _prepass.pkl):", bold=True, size=9.5)
table(["Reference","Meaning","Status"],
 [["ECMP","Equal-cost background routing (lower reference)","Available per-cycle ('emlu'); used as the ECMP background in the LP"],
  ["LP optimum / full-MCF","Per-cycle optimum used in PR","Available ('opt'); PR = optimum/achieved is reported throughout"],
  [f"{LABEL} (FIX1)","Proposed reward-gated method","Section 4 (mean PR 0.937-0.999, all PR>=0.90)"]], fs=8.5)
para("Baselines NOT re-run under the FIX1 model + protocol (do NOT copy Frozen Tier A numbers as FIX1):", bold=True, size=9.5)
table(["Baseline","FIX1 status"],
 [["OSPF / shortest-path","NOT RUN for FIX1 (see MISSING_BASELINES.md)"],
  ["Top-K demand","NOT RUN for FIX1"],
  ["Bottleneck-critical Top-K (standalone)","NOT RUN for FIX1"],
  ["GNN-only fixed K30 / LPD-only fixed K30","NOT RUN for FIX1"],
  ["GNN+LPD fixed K30 / K40 / K50","NOT RUN for FIX1"],
  ["GNN-only + reward gate / LPD-only + reward gate","NOT RUN for FIX1"]], fs=8.5)
note("These baseline rows are intentionally left as NOT RUN rather than filled with historical/frozen numbers. "
     "MISSING_BASELINES.md lists exactly what to regenerate. The report is not left without a baseline section, but "
     "no baseline number is fabricated.")

# 6
h1("6. K-Sensitivity / Gate vs Fixed-K")
para("A. Global fixed-K sweep (K10..K50) under the FIX1 model: NOT RUN for FIX1.", bold=True, size=9.5)
note("The global GNN+LPD fixed-K10..K50 sweep was not re-run on the FIX1 model, so a FIX1 global K-sensitivity table "
     "is not provided (would require copying frozen numbers — not done). See MISSING_BASELINES.md.")
para("B. VtlWavenet forced-K coverage frontier (historical: measured on the iter2 frozen model, NOT FIX1):", bold=True, size=9.5)
vs=pd.read_csv(RC/"FINAL_LEARNED_4OF5_ITER2_DDQN"/"vtl_scalability.csv")
table(["K budget","ODs opt.","Coverage %","Mean PR","Mean ms","P95 ms","p95<500?"],
  [[int(r.K_budget),int(r.ODs_optimized),f"{r.coverage_pct}%",f"{r.mean_PR:.4f}",f"{r.mean_ms:.0f}",f"{r.p95_ms:.0f}","yes" if r.under_500 else "NO"] for _,r in vs.iterrows()], fs=8.5)
note("Labeled HISTORICAL (iter2 frozen model). It demonstrates the structural coverage/runtime bound on VtlWavenet "
     "(8372 ODs): more coverage raises PR but only K<=500 stays within the p95 budget. FIX1 keeps VtlWavenet at its "
     "zero-shot steady state (mean PR 0.9373, 100% PR>=0.90).")

# 7
h1("7. Gate / RL Policy and Reward Settings")
para("Action distribution (per-cycle counts, FIX1):", bold=True, size=9.5)
rows=[]
for t in TOPO:
    d = adist.loc[t] if t in adist.index else pd.Series(dtype=float)
    rows.append([DISP[t]] + [int(d.get(a, 0)) for a in ACTS] + [M[t]["most_action"]])
table(["Topology"]+ACTS+["Most used"], rows, fs=8, c0=1.05)
note("Sticky / emergency actions: NOT used in FIX1 (the action space is exactly KEEP, K50, K100, K200, K300, K500, "
     "K800). KEEP = reuse previous routing; K* = optimize the top-K ranked ODs. The gate is the DDQN action policy; "
     "there is no RandomForest or sticky sub-gate in FIX1.")
para("Reward and training settings (FIX1):", bold=True, size=9.5)
table(["Parameter","Value"],
 [["Reward formula","r = 10*PR - 5*MLUex - 20*DB - 0.003*ms - 0.5*(k/nact); +10 if PR>=target; gated penalty if below; -25*((ms-500)/500) if ms>500"],
  ["lambda (DB weight W_DB)","20.0"],["MLU weight (W_MLU)","5.0"],["PR weight (W_PR)","10.0"],
  ["mu (decision-time weight W_MS)","0.003 (plus MS_GATE=25 over-budget penalty above 500 ms)"],
  ["target threshold","FlexDATE target where available, else 0.90 (policy never sees the target as input)"],
  ["gate model","Double-DQN QNet (MLP, 33-dim topology-agnostic state -> 7 actions)"],
  ["gamma","0.5"],["epsilon","0.07 (fixed during FIX1 fine-tune; not annealed)"],
  ["distillation weight (LAMBDA_DISTILL)","5.0 (MSE to frozen-teacher Q)"],
  ["fine-tune epochs","4"],["learning rate","2e-4"],
  ["TD updates / distill updates","13452 / 13452"],
  ["first-TM teacher-action preservation","ENABLED (cycle-0 action from frozen teacher; disclosed gated init)"],
  ["initialization","from frozen iter2 policy (not from scratch); full-data fine-tuned"]], fs=8)

# 8
h1("8. Zero-Shot Generalization")
rows=[]
for t in ["germany50","vtlwavenet2011"]:
    m=M[t]; rows.append([DISP[t],m["N"],f"{m['meanPR']:.4f}",f"{m['medPR']:.4f}",f"{m['pr95']:.0f}%",f"{m['pr90']:.0f}%",f"{m['minPR']:.4f}",f"{m['meanDB']:.4f}",f"{m['p95DB']:.4f}",f"{m['mean_ms']:.1f}",f"{m['p95_ms']:.1f}"])
zg=ev[ev.topology.isin(["germany50","vtlwavenet2011"])]
rows.append(["Combined zero-shot",len(zg),f"{zg.PR.mean():.4f}",f"{zg.PR.median():.4f}",f"{(zg.PR>=0.95).mean()*100:.0f}%",f"{(zg.PR>=0.90).mean()*100:.0f}%",f"{zg.PR.min():.4f}","-","-",f"{zg.decision_ms.mean():.1f}",f"{np.percentile(zg.decision_ms,95):.1f}"])
table(["Zero-shot topology","N","MeanPR","MedPR","PR>=.95","PR>=.90","MinPR","MeanDB","P95DB","mean ms","p95 ms"], rows, fs=8.5, c0=1.35)
note("Germany50 and VtlWavenet2011 are never trained on. VtlWavenet (8372 ODs) is coverage/runtime-bound (Section 6). "
     "Combined DB is left '-' because the two topologies' aggregate DB are reported separately (per-cycle DB not logged).")

# 9
h1("9. Decision-Time Breakdown")
para("Available split (2-way): a constant GNN inference offset per topology, and LP-solve + remaining pipeline = total - GNN offset.", bold=True, size=9.5)
rows=[]
for t in TOPO:
    rows.append([DISP[t],GNN_MS[t],f"{M[t]['mean_ms']-GNN_MS[t]:.1f}",f"{M[t]['mean_ms']:.1f}",f"{M[t]['p95_ms']:.1f}",f"{M[t]['max_ms']:.1f}"])
table(["Topology","GNN inference ms (offset)","LP-solve+rest mean ms","Total mean ms","Total p95 ms","Total max ms"], rows, fs=8.5, c0=1.1)
note("INSTRUMENTATION GAP: only total decision_ms and the constant GNN offset are recorded. The LPD-score, "
     "fusion/ranking, gate-decision, path-preparation, and DB-calculation components are NOT separately instrumented, "
     "so a per-component breakdown is not available without re-running with component timers. The 'LP-solve+rest' "
     "column lumps everything except the GNN offset.")
para("Decision time by action (FIX1, all topologies pooled):", bold=True, size=9.5)
rows=[]
for a in ACTS:
    ga=ev[ev.action==a]
    if len(ga): rows.append([a,len(ga),f"{ga.decision_ms.mean():.1f}",f"{np.percentile(ga.decision_ms,95):.1f}",f"{ga.decision_ms.max():.1f}"])
table(["Action","Rows","Mean ms","P95 ms","Max ms"], rows, fs=8.5)

# 10
h1("10. Solver / LP Problem Size / Hardware")
table(["Item","Value"],
 [["Solver","PuLP + CBC (COIN-OR)"],["Commercial solver?","No (open-source CBC)"],
  ["Solver status counted","Optimal (strict full-MCF numerator where available)"],
  ["Time limit","120 s (FIX1 eval) / 60 s (failure)"],["Threads","CBC default"],
  ["Hardware","Apple M3, 8 cores"],["OS","macOS 26.5 (Darwin 25.5.0, arm64)"],
  ["Python","3.12.12"],["Key packages","numpy 2.2.6, pandas 2.3.3, torch 2.10.0, pulp 3.3.0, networkx 3.5"]], fs=8.5)
para("LP problem size (selected-flow LP; variables ESTIMATED from selected-K x k_paths):", bold=True, size=9.5)
rows=[]
for t in TOPO:
    g=ev[ev.topology==t]; selK=g.selected_K.mean(); kp=np.where(g.selected_K.isin([500,800]),4,8).mean()
    est_vars=selK*kp
    rows.append([DISP[t],M[t]["nod"],f"{selK:.1f}",f"{kp:.1f}",f"~{est_vars:.0f}","not logged",f"{M[t]['mean_ms']:.1f}"])
table(["Topology","OD pairs","selected-K avg","k_paths avg","LP vars (est.)","LP constraints","Mean ms"], rows, fs=8.5, c0=1.1)
note("LP variables are ESTIMATED as selected-K x k_paths (per-path flow variables); exact variable/constraint counts "
     "are not logged by the solver wrapper, so they are labeled estimated. k_paths = 8 for K<=300 and 4 for K500/K800. "
     "Nodes/links per topology are not re-derived here; OD-pair counts are the actual ranked OD counts.")

# 11
h1("11. SDN / Mininet Operational Validation (retained simulation artifact)")
para("Decision time = AI + LP computation. Install time = OpenFlow rule installation. Recovery time = time after "
     "failure until traffic stabilizes. The table below is a RETAINED SDN-style operational validation artifact "
     "(Abilene + GEANT); it was NOT rerun on FIX1, and its own audit records LP-simulation mode because Mininet was "
     "not available on the macOS evaluation machine.", bold=True, size=9.5)
rows=[]
for _,r in sdn.iterrows():
    rows.append([DISP.get(r.topology,r.topology),r.scenario,f"{r.mean_throughput_mbps:.0f}",f"{r.mean_rtt_ms:.1f}",f"{r.mean_jitter_ms:.1f}",
        f"{r.mean_packet_loss_pct:.1f}",f"{r.mean_flow_rules:.0f}",f"{r.mean_install_ms:.2f}",f"{r.mean_recovery_ms:.1f}",int(r.disconnected_ODs),f"{r.mean_decision_ms:.1f}"])
table(["Topo","Scenario","Thr Mbps","RTT ms","Jitter ms","Loss %","Rules","Install ms","Recovery ms","Disc ODs","Decision ms"], rows, fs=7.5, c0=0.8)
note("RETAINED OPERATIONAL VALIDATION ARTIFACT; NOT RERUN ON FIX1. The decision-ms here is from the validation run, "
     "not FIX1's faster LP-eval. Audit scope: mode=simulate, mininet_available=false, topologies={Abilene, GEANT}, "
     "scenarios={normal, single_link_failure, two_link_failure, capacity_degradation_50, spike_x3, mixed_spike_failure}. "
     "Method label updated to RG-GNN-LPD, but this does not imply a fresh FIX1 Mininet rerun.")

# 12
h1("12. Failure Mission (FIX1)")
para("FIX1 periodic single-link failure robustness stress-test (full test streams):", bold=True, size=9.5)
rows=[]
for t in TOPO:
    g=pd.read_csv(FAIL/f"periodic_{t}.csv"); gf=g[g.is_failure==1]
    rows.append([DISP[t],len(g),int(g.is_failure.sum()),f"1/{NFAIL[t]}",f"{g.PR.mean():.4f}",f"{g.PR.min():.4f}",f"{(g.PR>=0.9).mean()*100:.0f}%",f"{gf.PR.mean():.4f}",f"{g.DB.mean():.4f}",int(g.disconnected_ODs.max())])
table(["Topology","TMs","#fail","every","MeanPR","MinPR","PR>=.90","PR(fail)","MeanDB","max disc ODs"], rows, fs=8)
para("Required wording: “This failure experiment is a robustness stress-test using full reroute over surviving "
     "paths. It is reported for resilience/PR behavior, not as the normal selected-flow runtime result.”", bold=True, size=9)
bullets([
 "Full reroute over surviving paths.",
 "Disconnected ODs are path-library limitations (the OD's precomputed candidate paths all traversed the failed link; "
 "NetworkX confirms the graph stays connected), NOT stale KEEP.",
 "Failure timing must NOT be mixed with the normal selected-flow runtime claim.",
 "Other scenarios (two-link, three-link, capacity-degradation-50, mixed spike+failure) were NOT re-run for FIX1 at "
 "the LP-eval level; they exist only in the retained SDN/Mininet artifact (Section 11). Frozen Tier A failure numbers "
 "are NOT used as FIX1."])
para("Disconnected-OD explanation (FIX1 stress-test):", bold=True, size=9.5)
table(["Topology","Max disc ODs","Connected after failure?","Cause"],
 [["Abilene","11","Yes (NetworkX)","Path-library limitation (Case B), not stale KEEP"],
  ["CERNET","40","Yes","Path-library limitation (Case B)"],
  ["VtlWavenet","6","Yes","Path-library limitation (Case B)"],
  ["Others","0","Yes","No disconnected ODs"]], fs=8.5)

# 13
h1("13. Figures (regenerated from FIX1 CSVs, RG-GNN-LPD)")
for nm,cap in [("fig1_pr_cdf.png","PR CDF"),("fig3_ms_cdf.png","Decision-time CDF"),("fig8_pr_vs_db.png","PR-DB tradeoff"),
               ("fig4_pr95.png","PR>=0.95 by topology"),("fig5_meandb.png","Mean DB by topology"),
               ("fig7_actiondist.png","Gate action distribution"),("fig10_fail_pr.png","Failure-cycle PR (stress-test)"),
               ("fig11_sdn.png","Retained SDN recovery (not rerun on FIX1)"),("fig_learn.png","FIX1 fine-tune curve")]:
    para(cap,bold=True,size=9,space=2,keep_next=True); fig(nm)
note("DB CDF and a normalized failure-MLU CDF are not shown: per-cycle realized DB is not logged, and raw MLU is not "
     "normalized across topologies (Section 4). Figures are regenerated from FIX1 CSVs; the retained SDN figure is "
     "labeled historical.")

# 14
h1("14. Reproducibility Checklist")
table(["Item","Value"],
 [["Train/eval command","/opt/homebrew/Caskroom/miniforge/base/bin/python3 scripts/phase1_5/run_fulldata_gated_fix1.py"],
  ["Report build","/opt/homebrew/Caskroom/miniforge/base/bin/python3 scripts/phase1_5/build_fix1_full_metrics.py then /opt/homebrew/Caskroom/miniforge/base/bin/python3 scripts/phase1_5/build_fix1_full_metrics_docx.py"],
  ["Failure run","/opt/homebrew/Caskroom/miniforge/base/bin/python3 scripts/phase1_5/run_periodic_failure_fix1.py"],
  ["Final model","FULLDATA_GATED_PRESERVED_FIX1/fulldata_gated_model.pt"],
  ["Dataset / prepass","results/.../condition_compliant_k10_k50/_prepass.pkl"],
  ["Topologies","abilene, geant, cernet, sprintlink, tiscali, ebone, germany50(zs), vtlwavenet2011(zs)"],
  ["Train/Val/Test split","Section 3 (Val = checkpoint-selection log, not a separate cycle count)"],
  ["k_paths","8 for K<=300, 4 for K500/K800"],["K actions","KEEP, K50, K100, K200, K300, K500, K800"],
  ["Training cap","full partitions (Abilene 2016, GEANT 672, others 200); 160-cap was the rejected Frozen Tier A"],
  ["alpha / db_budget","db_budget=0.051 normal; db_budget=1.0 gated cycle-0 (K>=300); db_weight=1e-6"],
  ["lambda/mu reward","W_DB=20, W_MS=0.003 (+MS_GATE=25); full weights in Section 7"],
  ["solver","PuLP + CBC; time limit 120 s"],["random seeds","42 (numpy/torch/random)"],
  ["CSV artifacts","fulldata_gated_eval_per_cycle.csv, normal_8topo_vs_target.csv, action_distribution.csv, flexdate_comparison.csv, PERIODIC_FAILURE_FIX1/periodic_*.csv"],
  ["hardware / OS","Apple M3 8-core / macOS 26.5 / Python 3.12.12"],
  ["Mininet / OVS version","Not available in retained artifact; sdn_method_audit.json records mode=simulate and mininet_available=false on macOS"]], fs=8)

# 15
h1("15. Claim Boundary")
bullets([
 f"Final learned model = FIX1 {LABEL}.",
 "Training = full-data fine-tuned.",
 "Learned FlexDATE wins = 3/5 (Abilene, CERNET, GEANT).",
 "4/5 only with the separate deployable Sprintlink K800 route (PR 0.9993, p95 ~325 ms).",
 "Sprintlink is NOT a learned win (learned Sprintlink PR = 0.9961 < 0.999).",
 "Tiscali remains a loss (PR 0.9525; target informal/non-source-locked).",
 "VtlWavenet is zero-shot and coverage/runtime-bound.",
 "The failure test is a robustness stress-test (periodic single-link); scenario-specific runtime-safe failure was NOT rerun for FIX1.",
 "Runtime on the largest topologies is borderline-near-500 ms, not a strict all-repetition guarantee.",
 "First-TM teacher-action preservation is part of the disclosed gated initialization mechanism."])

OUT = RPT / "FIX1_FullData_Gated_DDQN_Final_Report_FULL_METRICS.docx"
doc.save(str(OUT)); print("DOCX saved:", OUT)
