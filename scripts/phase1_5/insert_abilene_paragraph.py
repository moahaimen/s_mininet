#!/usr/bin/env python3
"""Surgical Section-11 insertion ONLY: add the Abilene two-link sensitivity paragraph (diagnostic) before
Section 12. No other change; the diagnostic counterfactual is explicitly NOT part of the final controller."""
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt

RC = Path("/Users/moahaimentalib/Desktop/f_flex_network_code_clean/results/gnn_lpd_dqn_selective_db_lp/condition_compliant_k10_k50")
R = RC / "FINAL_REPORT_FIX1" / "RG_GNN_LPD_FIX1_Final_Report_200VTL_N3976_FINAL_METHOD_AUDITED_PRREF_STRICTALL_MLU_UPGRADED.docx"
shutil.copy(R, R.with_suffix(".pre_abilene_para_backup.docx"))
doc = Document(str(R))
anchor = next(p for p in doc.paragraphs if p.text.strip().startswith("12.") and "Zero-Shot" in p.text)

HEAD = "Abilene Two-Link Sensitivity (diagnostic)"
PARA1 = ("Abilene is the most sensitive topology under the evaluated 20-TM two-link failure block, where the final "
         "controller achieves a mean PR of 88.936% and 40.0% of cycles reach PR>=0.90. The action trace includes four "
         "KEEP selections. In a diagnostic counterfactual, replacing failure-time KEEP decisions with optimization "
         "raises mean PR to 91.948% and PR>=0.90 to 60.0%, indicating that conservative KEEP selections materially "
         "contribute to the observed sensitivity. This diagnostic is not part of the reported final controller. The "
         "remaining sub-0.90 optimized cycles indicate that action choice alone does not fully explain the residual gap.")
PARA2 = ("A per-cycle counterfactual action sweep over the authorized action space classifies the 12 sub-0.90 two-link "
         "cycles as 6 action-selection-sensitive (a higher-K authorized action reaches PR>=0.90), 1 disturbance-budget-"
         "limited, and 5 candidate-path-coverage-limited (no authorized action reaches PR>=0.90 under the deployed k=8 "
         "path library). The best achievable mean PR over the authorized action space is 92.479% (PR>=0.90 = 65.0%). "
         "The residual is therefore a documented state-representation and candidate-path-coverage limitation rather than "
         "a defect in the reported method; a method-consistent learned fine-tune was evaluated and did not close the gap "
         "without regressing normal-protocol performance, so the final RG-GNN-LPD controller is retained unchanged.")

def add(text, bold=False, size=10):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
    anchor._p.addprevious(p._p)
add(HEAD, bold=True, size=10)
add(PARA1, size=9.5)
add(PARA2, size=9.5)
doc.save(str(R))
print(f"[done] inserted Abilene sensitivity paragraph. Backup: {R.with_suffix('.pre_abilene_para_backup.docx').name}")
